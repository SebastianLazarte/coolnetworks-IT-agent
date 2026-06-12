from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Remove compatibility mode restriction (Word 2010 compat = val 14)
settings = doc.settings.element
compat = settings.find(qn('w:compat'))
if compat is not None:
    for child in list(compat):
        if child.get(qn('w:name'), '') == 'compatibilityMode':
            compat.remove(child)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


# Title
title = doc.add_heading('Reporte del caso — Joachim Priedemann / Maswer', level=1)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Meta table
meta = [
    ('Cliente', 'Maswer'),
    ('Contacto principal', 'Vincenzo Valle'),
    ('Usuario final', 'Joachim Priedemann (oficina de Calden)'),
    ('Otros implicados', 'Oliver Orth, Carlos Vila Conde'),
    ('Tecnico asignado', 'Sebastian Lazarte Castellon (IT-Support-Germany / CoolNetworks)'),
    ('Fecha de apertura', '18-may-2026'),
    ('Estado actual', 'abierto - waiting on customer (disponibilidad sesion remota + envio equipo danado)'),
]
table = doc.add_table(rows=len(meta), cols=2)
table.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(4.5)

doc.add_paragraph()
add_hr(doc)

# Cronologia
doc.add_heading('Cronologia', level=2)
cronologia = [
    ('18-may', 'Vincenzo abre el ticket: (A) configurar portatil nuevo para Joachim, (B) problema de conexion con Starkis en Calden.'),
    ('18-may 21:00', 'Respuesta inicial - acuse de recibo y solicitud de informacion (Miguel Rubira).'),
    ('18-19 may', 'Triage interno: se separan los dos asuntos. Intento de videollamada con Joachim - no se conecto. Comprobaciones de calidad de conexion desde nuestro lado - todo en orden. Llamada a la oficina de Calden - no se pudo contactar.'),
    ('19-may', 'Vincenzo aporta info: portatil = HP EliteBook 840 G9, movil de Joachim +49 170 3737610. Aclara que Joachim no puede trabajar por Starkis.'),
    ('19-may', 'Llamada a las 16:30 acordada con Vincenzo. Reformula: portatil actual se apaga cada ~10 min (bloqueo real); Starkis es problema separado y antiguo.'),
    ('19-may', 'Instalacion del nuevo portatil ejecutada. Se deshabilita temporalmente el 2FA y se reactiva. Ticket cerrado con confirmacion.'),
    ('22-may', 'Cliente reabre (ticket nuevo): portatil no esta completamente listo, debe tener la misma configuracion y software. Asumen que tenemos backup.'),
    ('22-28 may', 'Se aclara: no existe backup y el portatil viejo es demasiado inestable para clonarlo. Se propone reinstalacion estandar + lista de software/accesos facilitada por Joachim. Mensaje traducido al aleman.'),
    ('29-may', 'Respuesta de Oliver: portatil se reenvía a Hennef, ira en persona y propone llevar a Sebastian. Confirma solo perfil local Win11, sin acceso al servidor. Joachim tiene 3 dias antes de vacaciones.'),
    ('29-may', 'Constraint del tecnico: post-operacion de tendon de Aquiles - enfoque mixto: Oliver en sitio + Sebastian remoto por AnyDesk. A la espera de fecha.'),
    ('03-jun', 'Vincenzo pregunta por que se solicitan los equipos. Se le explica la situacion.'),
    ('03-jun', 'Se propone a Vincenzo dos opciones: sesion remota (con alguien en sitio) o envio de ambos equipos a Espana.'),
    ('03-jun', 'Decision: sesion remota para configurar el nuevo HP EliteBook 840 G9 (domain join, perfil, acceso al servidor). Portatil danado se envia a Espana por protocolo contractual. Pendiente confirmar hora con Vincenzo.'),
]
cron_table = doc.add_table(rows=1 + len(cronologia), cols=2)
cron_table.style = 'Table Grid'
for cell, text in zip(cron_table.rows[0].cells, ['Fecha', 'Evento']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1F497D')
for i, (fecha, evento) in enumerate(cronologia):
    row = cron_table.rows[i + 1]
    row.cells[0].text = fecha
    row.cells[1].text = evento
    row.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0:
        shade_cell(row.cells[0], 'EFF3F9')
        shade_cell(row.cells[1], 'EFF3F9')
cron_table.columns[0].width = Inches(1.3)
cron_table.columns[1].width = Inches(5.2)

doc.add_paragraph()
add_hr(doc)

# Trabajo realizado
doc.add_heading('Trabajo realizado', level=2)
subsections = [
    ('Comunicacion con cliente', [
        '1 acuse de recibo con preguntas de scoping.',
        '1 update de estado con resultados de las pruebas.',
        '1 mensaje confirmando reunion 16:30.',
        '1 mensaje de resolucion (cierre del primer ticket) explicando la deshabilitacion temporal del 2FA.',
        '4 iteraciones tras la reapertura (sin backup, no se puede clonar, peticion de lista de software/AnyDesk, respuesta a Oliver).',
        'Aclaracion a Vincenzo sobre el motivo de solicitar los equipos.',
        'Coordinacion del plan final: sesion remota + envio del equipo danado a Espana por protocolo contractual.',
    ]),
    ('Diagnostico y pruebas tecnicas', [
        'Comprobaciones de calidad de conexion desde nuestro lado - todo OK (Starkis no es problema de infraestructura nuestra).',
        'Identificacion del fallo del portatil viejo (apagado cada ~10 min).',
        'Constatacion de que el clonado 1:1 no es viable (sin backup + maquina origen no estable).',
    ]),
    ('Decisiones de triage', [
        'Separacion de los dos asuntos en tickets independientes.',
        'Priorizacion: portatil = P2 (usuario bloqueado), Starkis = P3 (no bloqueante, problema antiguo, un solo usuario).',
        'Sin escalado: nada apunto a Cybersecurity ni N2 Systems.',
    ]),
    ('Setup ejecutado en el primer ciclo', [
        'Imagen del HP EliteBook 840 G9, 2FA gestionado de forma segura, devolucion prevista a Joachim. Posteriormente el cliente reporto que faltaba domain join + perfil de usuario - de ahi la reapertura.',
    ]),
]
for heading, items in subsections:
    doc.add_heading(heading, level=3)
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Tiempo
doc.add_heading('Tiempo', level=2)
tiempos = [
    ('Elapsed total', '16 dias (18-may - 03-jun), todavia abierto.'),
    ('Primer ciclo', '~24-36 h desde apertura hasta resolucion reportada.'),
    ('Segundo ciclo (reapertura)', '7 dias y contando, bloqueado en logistica y disponibilidad del cliente.'),
    ('Tiempo activo del tecnico (estimado)', '4-6 h efectivas: triage, intentos de contacto, llamada 16:30, setup 2FA, redaccion y traduccion de comunicaciones.'),
]
for label, value in tiempos:
    p = doc.add_paragraph()
    p.add_run(label + ': ').bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
add_hr(doc)

# Estado actual
doc.add_heading('Estado actual y pendientes', level=2)
p = doc.add_paragraph()
p.add_run('Estado: ').bold = True
run = p.add_run('abierto - waiting on customer')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)

doc.add_heading('Para cerrar el portatil necesitamos:', level=3)
pendientes = [
    'Confirmacion de disponibilidad y hora de Vincenzo para la sesion remota.',
    'Alguien en sitio (Calden/Hennef) durante la sesion para gestionar el equipo.',
    'Lista de software y accesos de Joachim.',
    'Creacion de cuenta AD / perfil de dominio antes de la sesion.',
    'Envio del portatil danado a Espana (protocolo contractual).',
]
for item in pendientes:
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pendiente posterior: ').bold = True
p.add_run('verificacion de ')
p.add_run('Starkis').bold = True
p.add_run(' sobre el portatil nuevo una vez Joachim trabaje con el (Ticket B, P3).')

out = r'g:\Mi unidad\Emprendimientos\AI_Automation_Agency\AIClieffNotes\WeeklyCompetions\week3\coolnetworks-tecnico-n1\coolnetworks-IT-Agent\reports\2026-05-29_maswer_joachim-priedemann.docx'
doc.save(out)

# Remove Windows Zone.Identifier ADS so Word doesn't open it in Protected View
import subprocess
subprocess.run(
    ['powershell', '-Command', f'Unblock-File -LiteralPath "{out}"'],
    check=True
)
print('OK')
