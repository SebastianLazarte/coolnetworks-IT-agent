from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

settings = doc.settings.element
compat = settings.find(qn('w:compat'))
if compat is not None:
    for child in list(compat):
        if child.get(qn('w:name'), '') == 'compatibilityMode':
            compat.remove(child)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


# Title
title = doc.add_heading('Reporte del caso - Joaquin / Reinstalacion de STAkis Profi', level=1)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Meta table
meta = [
    ('Cliente', '[a confirmar - cliente de STAHLGRUBER]'),
    ('Usuario final', 'Joaquin [apellido a confirmar]'),
    ('Otros implicados', 'Oliver Orth (tecnico con experiencia previa instalando STAkis)'),
    ('Tecnico asignado', 'Sebastian Lazarte (CoolNetworks)'),
    ('Fecha de apertura', '05-jun-2026'),
    ('Estado actual', 'en investigacion - pendiente respuesta de Oliver'),
]
table = doc.add_table(rows=len(meta), cols=2)
table.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(4.5)

doc.add_paragraph()
add_hr(doc)

# Contexto
doc.add_heading('Contexto', level=2)
ctx = doc.add_paragraph()
ctx.add_run(
    'Se solicita instalar STAkis Profi (software de gestion de talleres y comercio de '
    'repuestos de STAHLGRUBER) en el equipo nuevo de Joaquin. El usuario ya lo tenia '
    'instalado en su maquina anterior, por lo que se trata de una reinstalacion tras '
    'cambio de equipo, no de un alta nueva.'
)
ctx2 = doc.add_paragraph()
ctx2.add_run(
    'STAkis Profi no es software de descarga publica: se distribuye unicamente a traves '
    'del portal de clientes de STAHLGRUBER (kunden.stahlgruber.de) o por contacto directo '
    'con stakis.support@stahlgruber.de. Requiere licencia/credenciales del cliente.'
)

doc.add_paragraph()
add_hr(doc)

# Cronologia
doc.add_heading('Cronologia', level=2)
cronologia = [
    ('05-jun', 'Solicitud inicial del usuario: "quiero instalar starkis". Aclaracion con captura del producto -> confirmado STAkis Profi (STAHLGRUBER).'),
    ('05-jun', 'Investigacion de la via oficial de instalacion: portal de clientes STAHLGRUBER, hotline 0800 5782-547, mail stakis.support@stahlgruber.de. Descartado uso de copias no oficiales (mhhauto, obd2diagnoseshop, etc.) por riesgo de malware y problemas de licencia.'),
    ('05-jun', 'Confirmado que el cliente ya es cliente de STAHLGRUBER -> via oficial viable.'),
    ('05-jun', 'Redactado correo en aleman dirigido al soporte de STAkis (stakis.support@stahlgruber.de) para solicitar instalador, licencia y requisitos de sistema. No enviado todavia - pendiente decision interna.'),
    ('05-jun', 'Cambio de enfoque: se prioriza consultar primero a Oliver Orth, que ya tiene experiencia instalando STAkis, antes de escalar al soporte del proveedor.'),
    ('05-jun', 'Redactado correo en ingles a Oliver pidiendole su procedimiento habitual de instalacion (origen del instalador, licencia/activacion, problemas comunes). Pendiente envio y respuesta.'),
]
cron_table = doc.add_table(rows=1 + len(cronologia), cols=2)
cron_table.style = 'Table Grid'
for cell, text in zip(cron_table.rows[0].cells, ['Fecha', 'Evento']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1F497D')
for i, (fecha, evento) in enumerate(cronologia):
    row = cron_table.rows[i + 1]
    row.cells[0].text = fecha
    row.cells[1].text = evento
    row.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0:
        shade_cell(row.cells[0], 'EFF3F9')
        shade_cell(row.cells[1], 'EFF3F9')
cron_table.columns[0].width = Inches(1.3)
cron_table.columns[1].width = Inches(5.2)

doc.add_paragraph()
add_hr(doc)

# Trabajo realizado
doc.add_heading('Trabajo realizado', level=2)
subsections = [
    ('Investigacion', [
        'Identificacion del software a partir de la captura del usuario (no era "Starlink" ni "Starship" - era STAkis Profi de STAHLGRUBER).',
        'Localizacion de los canales oficiales de distribucion y soporte: portal kunden.stahlgruber.de, hotline 0800 5782-547, mail stakis.support@stahlgruber.de, PDF de producto kunden.stahlgruber.de/downloads/STAkis_Profi_web.pdf.',
        'Revision de stahlgruber.de/en/services/digital-customer-systems para confirmar canales de soporte.',
        'Descartadas fuentes no oficiales por riesgo de seguridad y legal.',
    ]),
    ('Comunicacion (redactada, pendiente de envio)', [
        'Correo a soporte STAkis (aleman): solicita instalador, licencia y requisitos para el equipo destino. Incluye plantilla con campos a completar (nro cliente, SO, single/multi-puesto, instalacion previa).',
        'Correo a Oliver Orth (ingles): consulta su procedimiento habitual - origen del instalador, licenciamiento, errores tipicos. Tono cordial, ofrece llamada si es mas practico.',
    ]),
    ('Decisiones de triage', [
        'Categoria: Software (3rd-party / STAHLGRUBER).',
        'Prioridad: P3 (solicitud estandar, sin impacto operativo inmediato - Joaquin puede seguir trabajando en otras tareas).',
        'Via elegida: consultar primero al tecnico interno con experiencia (Oliver) antes de abrir caso con el proveedor.',
    ]),
]
for heading, items in subsections:
    doc.add_heading(heading, level=3)
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Estado actual
doc.add_heading('Estado actual y pendientes', level=2)
p = doc.add_paragraph()
p.add_run('Estado: ').bold = True
run = p.add_run('en investigacion - pendiente respuesta de Oliver')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)

doc.add_heading('Proximos pasos:', level=3)
pendientes = [
    'Enviar el correo a Oliver con la consulta del procedimiento.',
    'Segun la respuesta de Oliver: si tiene instalador + licencia documentados internamente -> proceder con la instalacion remota; si no -> enviar el correo en aleman al soporte de STAkis pidiendo instalador y credenciales del cliente.',
    'Confirmar con Joaquin / cliente: numero de cliente STAHLGRUBER, modalidad de la instalacion previa (single-puesto o multi-puesto), uso conjunto con STAkis-S u otros modulos.',
    'Coordinar ventana de instalacion remota con Joaquin (AnyDesk / Quick Assist / herramienta interna definida).',
]
for item in pendientes:
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Notas: ').bold = True
p.add_run(
    'Sin actividad sobre la maquina de Joaquin todavia - toda la sesion fue investigacion '
    'y preparacion de comunicaciones. No se ha gastado tiempo del usuario final mas alla '
    'de la captura inicial del producto.'
)

out = r'g:\Mi unidad\Emprendimientos\AI_Automation_Agency\AIClieffNotes\WeeklyCompetions\week3\coolnetworks-tecnico-n1\coolnetworks-IT-Agent\reports\2026-06-05_joaquin_stakis-reinstall.docx'
doc.save(out)

import subprocess
subprocess.run(
    ['powershell', '-Command', f'Unblock-File -LiteralPath "{out}"'],
    check=True
)
print('OK')
