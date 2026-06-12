# -*- coding: utf-8 -*-
"""Genera el informe ejecutivo en formato Word (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paleta
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x6D, 0xA4)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color

def add_heading(text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        # borde inferior
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E6DA4")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_bold_lead(lead, rest):
    p = doc.add_paragraph()
    r = p.add_run(lead)
    r.bold = True
    p.add_run(rest)
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def add_check(text):
    p = doc.add_paragraph()
    r = p.add_run("✓  ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x8B, 0x2E)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def styled_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], "1F3A5F")
        set_cell_text(hdr[i], h, bold=True, white=True, size=10)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
        if ridx % 2 == 1:
            for c in cells:
                shade_cell(c, "EAF1F8")
    return table

# ---------- PORTADA / TÍTULO ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("INFORME EJECUTIVO")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
r = sub.add_run("Análisis de las incidencias IT en Maswer Alemania")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run("Periodo analizado: Mayo – Junio 2026")
r.font.size = Pt(11)
r.font.color.rgb = GREY
meta2 = doc.add_paragraph()
r = meta2.add_run("Preparado por CoolNetworks  ·  11 de junio de 2026")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---------- RESUMEN EJECUTIVO ----------
add_heading("Resumen ejecutivo", 1)
add_para("Tras analizar la documentación disponible, los tickets registrados y las actuaciones realizadas por CoolNetworks en Alemania durante los meses de mayo y junio de 2026, no se observa una situación de mal funcionamiento generalizado del servicio IT.")
add_para("La mayoría de las incidencias registradas corresponden a actividades habituales de soporte técnico: administración de usuarios, gestión de permisos, configuración de equipos y asistencia funcional a usuarios.")
add_para("Únicamente se identifica una incidencia crítica de infraestructura que provocó una pérdida temporal de conectividad y acceso a recursos corporativos. Dicha incidencia fue causada por un equipo ubicado en un centro de datos externo y fue restablecida el mismo día tras la intervención del proveedor responsable.")
add_para("No se ha encontrado evidencia de pérdida de datos, brechas de seguridad, incumplimientos reiterados de SLA ni degradación estructural de los servicios IT en Alemania.")

# ---------- PRINCIPALES INCIDENCIAS ----------
add_heading("Principales incidencias detectadas", 1)

add_heading("1. Problema con Microsoft Visio", 2)
add_bold_lead("Usuario afectado: ", "Maurizio Carroccia")
add_para("Inicialmente se reportó una incidencia relacionada con la aplicación Visio. Tras el análisis se comprobó que el equipo del usuario tenía Office instalado pero no la aplicación Microsoft Visio, que era lo que faltaba. Se verificó la licencia asignada y se completó la instalación, dejándola funcional para el usuario.")
add_bold_lead("Conclusión: ", "no se trataba de una caída del servicio ni de un fallo de infraestructura, sino de una necesidad de instalación de software en el equipo.")

add_heading("2. Incidencia de conectividad WiFi", 2)
add_para("Durante el periodo se registró una incidencia de conectividad. La actuación de CoolNetworks fue inmediata y la propia usuaria confirmó la recuperación del servicio en menos de una hora desde la apertura del ticket.")
add_bold_lead("Conclusión: ", "incidencia puntual resuelta satisfactoriamente, sin impacto prolongado.")

add_heading("3. Caída general de conectividad", 2)
add_para("Esta ha sido la incidencia más relevante registrada durante el periodo.")
add_bold_lead("Servicios afectados:", "")
add_bullet("Internet")
add_bullet("VPN corporativa")
add_bullet("Carpetas compartidas")
add_bullet("Recursos corporativos centralizados")
add_bold_lead("Causa identificada: ", "el equipo que conectaba la oficina con la infraestructura corporativa, ubicado en un centro de datos externo, dejó de responder. El proveedor responsable restauró dicho equipo y, posteriormente, CoolNetworks verificó el correcto funcionamiento de todos los servicios.")
add_bold_lead("Impacto: ", "paralización temporal de determinadas tareas operativas durante varias horas.")
add_bold_lead("Resultado:", "")
add_bullet("Recuperación del servicio el mismo día.")
add_bullet("Sin pérdida de información.")
add_bullet("Sin incidentes de ciberseguridad.")
add_bullet("Comunicación continua con los usuarios afectados.")

add_heading("4. Lentitud en las actualizaciones (Vaihingen)", 2)
add_para("Se reportaron problemas de rendimiento durante procesos de actualización. Las investigaciones apuntaron a la necesidad de revisar:")
add_bullet("Operador de comunicaciones.")
add_bullet("Ancho de banda disponible.")
add_bullet("Equipamiento local.")
add_bullet("Infraestructura WAN.")
add_para("No existe evidencia técnica que permita atribuir esta situación a una actuación incorrecta de CoolNetworks.")
add_bold_lead("Conclusión: ", "incidencia de rendimiento en fase de diagnóstico, con múltiples factores externos involucrados.")

add_heading("5. Intento de phishing (ticket 641)", 2)
add_para("Se recibió un correo que indicaba falsamente que el almacenamiento del buzón estaba lleno y que en 22 días se dejaría de recibir correo. El mensaje llegó a una cuenta que en la práctica no se utiliza, lo que descartaba un aviso legítimo de cuota. El caso se gestionó a través del área de Ciberseguridad. Tras la revisión:")
add_bullet("Se identificó el mensaje como phishing (intento de generar urgencia para robar credenciales).")
add_bullet("El usuario no introdujo credenciales.")
add_bullet("No se produjo acceso no autorizado.")
add_bold_lead("Conclusión: ", "incidente de seguridad detectado y contenido correctamente, sin consecuencias para la organización.")

# ---------- GESTIÓN DE PERMISOS ----------
add_heading("Gestión de permisos y accesos corporativos", 1)
add_para("Uno de los aspectos más relevantes detectados es que una parte importante de los tickets no corresponden a averías o fallos de servicio, sino a solicitudes de acceso a recursos corporativos.")

add_heading("Carpeta HR", 2)
add_para("Se solicitaron accesos a la carpeta de RR.HH. para dos usuarios. El proceso seguido fue:")
doc.add_paragraph("Solicitud formal.", style="List Number")
doc.add_paragraph("Validación de autorizaciones (la carpeta HR contiene datos sensibles).", style="List Number")
doc.add_paragraph("Asignación de permisos.", style="List Number")
doc.add_paragraph("Confirmación por parte de los usuarios.", style="List Number")
add_para("La gestión se cerró con la confirmación de acceso por parte de los propios usuarios.")
add_bold_lead("Punto importante: ", "la carpeta HR está especialmente protegida y su control final lo tiene el proveedor externo conet.de. Por eso no pudimos verificar los permisos directamente desde el servidor: nuestra cuenta no tiene el acceso necesario, y es conet.de quien debe concederlo. El acceso quedó confirmado por los propios usuarios, pero conviene revisar todos los accesos a estas carpetas para asegurar que cada usuario tiene el permiso que le corresponde.")

add_heading("Acceso a aplicación financiera (Profi Cash)", 2)
add_para("Solicitud de acceso a la herramienta Profi Cash para un usuario. Se identificó el entorno de acceso, se configuró del lado del cliente y se remitió el acceso al usuario para su uso. No requirió instalación de software adicional en el equipo.")

add_heading("Alta de nuevo usuario", 2)
add_para("Solicitud de creación de una cuenta nueva con los mismos accesos que una empleada existente. Se creó la cuenta, se replicaron los permisos y se asignó la licencia y el buzón correspondientes. Actuación dentro de la administración normal de usuarios y sistemas corporativos.")

# ---------- TIEMPOS DE RESPUESTA ----------
add_heading("Tiempos de respuesta observados", 1)
add_para("Los tiempos de respuesta obtenidos a partir de los tickets revisados muestran una actuación rápida y alineada con las buenas prácticas de soporte IT.")
styled_table(
    ["Incidencia", "Tiempo de respuesta"],
    [
        ["Problema Visio", "11 minutos"],
        ["Incidencia WiFi", "Atención inmediata"],
        ["Caída general de conectividad", "37 minutos"],
        ["Solicitud acceso HR", "42 minutos"],
        ["Incidente de phishing", "Mismo día"],
    ],
)
doc.add_paragraph()
add_para("Las incidencias críticas fueron atendidas en menos de una hora y las solicitudes administrativas se gestionaron en plazos razonables.")

# ---------- DISTRIBUCIÓN ----------
add_heading("Distribución real de las actuaciones", 1)
add_para("El análisis de los tickets permite estimar la siguiente distribución del trabajo:")
styled_table(
    ["Tipo de actuación", "Peso aproximado"],
    [
        ["Gestión de permisos y accesos", "35 %"],
        ["Soporte a usuarios y software", "25 %"],
        ["Configuración de equipos", "15 %"],
        ["Incidencias de conectividad", "15 %"],
        ["Seguridad y consultas", "10 %"],
    ],
)
doc.add_paragraph()
add_para("Esta distribución demuestra que gran parte de la actividad corresponde a soporte operativo normal y no a incidencias derivadas de fallos de infraestructura.")

# ---------- VALORACIÓN ----------
add_heading("Valoración del servicio", 1)
add_para("La documentación revisada acredita que:")
add_check("Los tickets fueron atendidos.")
add_check("Se mantuvo comunicación constante con los usuarios.")
add_check("Las incidencias críticas fueron escaladas adecuadamente.")
add_check("Se realizaron actuaciones remotas de soporte.")
add_check("Se gestionaron correctamente los accesos y permisos solicitados.")
add_check("Los problemas reportados fueron resueltos o correctamente diagnosticados.")
add_para("No se observa evidencia de incumplimientos reiterados ni de degradación estructural del servicio IT en Alemania.")

# ---------- RECOMENDACIONES ----------
add_heading("Recomendaciones de mejora", 1)
add_para("Con independencia de que el servicio se haya prestado correctamente, se recomienda:")
add_bold_lead("1. Mejora de la documentación técnica. ", "Mantener evidencias periódicas de firewalls, VPN, conectividad WAN y rendimiento de enlaces.")
add_bold_lead("2. Informes mensuales de servicio. ", "Presentar un informe mensual que incluya: tickets abiertos, tickets cerrados, tiempo medio de respuesta, incidencias críticas y acciones preventivas.")
add_bold_lead("3. Redundancia de comunicaciones. ", "Evaluar soluciones de respaldo para minimizar el impacto de incidencias en proveedores externos o centros de datos.")
add_bold_lead("4. Revisión de accesos a carpetas. ", "Comprobar todos los accesos a las carpetas corporativas para confirmar que cada usuario tiene exactamente el permiso que le corresponde. Para poder hacerlo, conet.de debe conceder a la cuenta de CoolNetworks el acceso necesario para revisar y verificar los permisos a nivel de servidor; sin ese acceso, quedan puntos ciegos fuera de nuestro alcance.")

# Pie
doc.add_paragraph()
foot = doc.add_paragraph()
r = foot.add_run("Documento preparado por CoolNetworks · 11 de junio de 2026.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

out = r"g:\Mi unidad\Emprendimientos\AI_Automation_Agency\AIClieffNotes\WeeklyCompetions\week3\coolnetworks-tecnico-n1\coolnetworks-IT-Agent\reports\2026-06-11_Informe-Ejecutivo-Maswer.docx"
doc.save(out)
print("Saved:", out)
