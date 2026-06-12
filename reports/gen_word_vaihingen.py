from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

settings = doc.settings.element
compat = settings.find(qn('w:compat'))
if compat is not None:
    for child in list(compat):
        if child.get(qn('w:name'), '') == 'compatibilityMode':
            compat.remove(child)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


# Title
title = doc.add_heading('Reporte de caso — Maswer GmbH (Vaihingen) / Caída total de internet — túnel S2S_Vaihingen', level=1)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Meta table
meta = [
    ('Cliente', 'Maswer GmbH (Maswer Alemania — Vaihingen an der Enz)'),
    ('Contacto', 'Alketa Vrella'),
    ('Infraestructura clave', 'Sophos XGS (admin@XGSDEVAI01, serial X125058C8GJ8H51, vía Sophos Central) · túnel IPsec S2S_Vaihingen (route-based, perfil S2SVaihingen)'),
    ('Técnico asignado', 'Sebastian Lazarte Castellón (CoolNetworks)'),
    ('Fecha', '08-Jun-2026'),
    ('Estado actual', 'Resuelto (sitio online) · pendientes de prevención abiertos para N2 Sistemas'),
]
table = doc.add_table(rows=len(meta), cols=2)
table.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
table.columns[0].width = Inches(1.8)
table.columns[1].width = Inches(4.7)

doc.add_paragraph()
add_hr(doc)

# Triage
doc.add_heading('Triage del ticket', level=2)
triage = [
    ('Categoría', 'Red e Infraestructura'),
    ('Prioridad', 'P1 - Crítica (caída total de sitio, negocio parado)'),
    ('Grupo', 'N1 primer contacto → N2 Sistemas'),
    ('SLA inicial', 'Respuesta en 15 min · resolución objetivo 4 h'),
    ('Idioma del ticket', 'Alemán → respuesta en inglés (no inventamos alemán)'),
    ('Empresa en Freshdesk', 'Maswer Alemania (confirmar mapeo del contacto)'),
]
t_table = doc.add_table(rows=1 + len(triage), cols=2)
t_table.style = 'Table Grid'
for cell, text in zip(t_table.rows[0].cells, ['Campo', 'Valor']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1F497D')
for i, (k, v) in enumerate(triage):
    row = t_table.rows[i + 1]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0:
        shade_cell(row.cells[0], 'EFF3F9')
        shade_cell(row.cells[1], 'EFF3F9')

doc.add_paragraph()
add_hr(doc)

# Cómo llegó el caso
doc.add_heading('Cómo llegó el caso', level=2)
intake = [
    'Primer aviso por WhatsApp (canal fuera de plataforma). Alketa indica que cree haber abierto un ticket previo y que el Wi-Fi está caído desde por la mañana, no pueden trabajar.',
    'Segundo mensaje, ya en el ticket, en alemán. Aporta dato clave nuevo: no hay internet ni por Wi-Fi ni por cable LAN. Solicita explícitamente envío de técnico in situ.',
]
for item in intake:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().add_run('Acción de N1:').bold = True
n1_actions = [
    'Consolidar en un único ticket (sin abrir duplicados).',
    'Redirigir educadamente desde WhatsApp al portal de soporte (donde están las herramientas de diagnóstico).',
    'Confirmar que el contacto está mapeado a la compañía Maswer Alemania.',
]
for item in n1_actions:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Diagnóstico inicial
doc.add_heading('Diagnóstico inicial', level=2)
diag = [
    'Caída total de sitio, LAN + Wi-Fi simultáneamente → descarta fallo solo de Wi-Fi (APs/controladora).',
    'Apunta a enlace WAN/ISP o firewall perimetral / routing.',
    'El firewall Sophos XGS sigue alcanzable vía Sophos Central → el dispositivo está vivo, el problema está en el camino del tráfico, no en el equipo.',
]
for item in diag:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Hallazgo Sophos
doc.add_heading('Hallazgo en Sophos XGS (Site-to-site VPN → IPsec)', level=2)
doc.add_paragraph('Una única conexión configurada:')
hallazgo = [
    ('Nombre', 'S2S_Vaihingen'),
    ('Perfil', 'S2SVaihingen'),
    ('Tipo', 'Route-based (Tunnel interface)'),
    ('Active', 'Verde (habilitado administrativamente)'),
    ('Connection', 'ROJO — túnel caído, no establecido'),
    ('Failover group', 'Sin registros'),
]
h_table = doc.add_table(rows=1 + len(hallazgo), cols=2)
h_table.style = 'Table Grid'
for cell, text in zip(h_table.rows[0].cells, ['Campo', 'Valor']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1F497D')
for i, (k, v) in enumerate(hallazgo):
    row = h_table.rows[i + 1]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0:
        shade_cell(row.cells[0], 'EFF3F9')
        shade_cell(row.cells[1], 'EFF3F9')

p = doc.add_paragraph()
p.add_run('Causa raíz muy probable: ').bold = True
p.add_run('el sitio enruta su tráfico/internet a través del túnel S2S_Vaihingen. Túnel caído + sin grupo de failover = sitio entero sin internet por LAN y Wi-Fi. Encaja al 100% con los síntomas.')

doc.add_paragraph()
add_hr(doc)

# Resolución
doc.add_heading('Resolución', level=2)
res = [
    'El IT anterior del cliente, Oliver, restauró el túnel desde su lado.',
    'Verificado: Connection en verde, internet de vuelta en Maswer Vaihingen.',
    'No se ejecutaron cambios en producción desde CoolNetworks durante este incidente.',
]
for item in res:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Comunicaciones
doc.add_heading('Comunicaciones enviadas', level=2)
com = [
    'Acuse inicial al cliente (inglés, en el ticket). Confirmación de criticidad, consolidación con el aviso de WhatsApp, preguntas operativas mínimas (luces del router/firewall, alcance).',
    'Mensaje en WhatsApp. Redirección proactiva al ticket, sin preguntar permiso: instrucciones claras (abrir el ticket, mirar la respuesta, enviar las fotos) y oferta de llamada como respaldo.',
    'Email de agradecimiento a Oliver (no técnico). Breve: gracias por resolverlo, confirmación de que queda documentado por nuestra parte.',
]
for i, item in enumerate(com, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Línea de tiempo
doc.add_heading('Línea de tiempo', level=2)
cronologia = [
    ('08-Jun, mañana', 'Inicio de la caída en Maswer Vaihingen (LAN + Wi-Fi).'),
    ('08-Jun', 'Alketa contacta por WhatsApp pidiendo ayuda urgente.'),
    ('08-Jun', 'Llega mensaje en alemán al ticket: confirma caída de LAN además de Wi-Fi; pide técnico in situ.'),
    ('08-Jun', 'N1 clasifica como P1, redirige a la plataforma, prepara escalado a N2 Sistemas.'),
    ('08-Jun', 'Revisión del Sophos XGS → S2S_Vaihingen con Connection en rojo identificado como causa raíz probable.'),
    ('08-Jun', 'Oliver (IT previo del cliente) restablece el túnel. Sitio recupera internet.'),
    ('08-Jun', 'Email de agradecimiento a Oliver. Runbook documentado en memoria del agente.'),
]
cron_table = doc.add_table(rows=1 + len(cronologia), cols=2)
cron_table.style = 'Table Grid'
for cell, text in zip(cron_table.rows[0].cells, ['Fecha', 'Evento']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1F497D')
for i, (fecha, evento) in enumerate(cronologia):
    row = cron_table.rows[i + 1]
    row.cells[0].text = fecha
    row.cells[1].text = evento
    row.cells[0].paragraphs[0].runs[0].bold = True
    if i % 2 == 0:
        shade_cell(row.cells[0], 'EFF3F9')
        shade_cell(row.cells[1], 'EFF3F9')
cron_table.columns[0].width = Inches(1.5)
cron_table.columns[1].width = Inches(5.0)

doc.add_paragraph()
add_hr(doc)

# Pendientes
doc.add_heading('Pendientes (prevención — para N2 Sistemas)', level=2)
doc.add_paragraph('El "bounce" del túnel es primeros auxilios, no la cura. Si no se cierran estos puntos, el incidente se va a repetir:')
pendientes = [
    'Configurar un Failover group en el Sophos para el sitio de Vaihingen. Sin camino de respaldo, una sola caída del túnel = sitio entero a oscuras.',
    'Alerta de tunnel-down en Sophos Central o en el RMM. Hoy el monitor es el cliente.',
    'Análisis de causa raíz del túnel: revisar el perfil S2SVaihingen (IKE, lifetimes, DPD, rekey); estabilidad de la línea ISP en ambos extremos; IP WAN dinámica en el peer sin DDNS; fragmentación/MTU.',
    'Confirmar la ruta por defecto del sitio: ¿realmente todo el tráfico de Vaihingen sale por este túnel? Documentarlo.',
    'Definir si WhatsApp es canal de soporte acordado para Maswer Alemania. Si no lo es, dejarlo por escrito para no normalizar avisos P1 fuera de plataforma.',
]
for item in pendientes:
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Runbook
doc.add_heading('Runbook de recuperación (si vuelve a caer)', level=2)
doc.add_paragraph('Guardado en memoria del agente: maswer-vaihingen-s2s-vpn-runbook.md.')
doc.add_paragraph('Resumen operativo:')
runbook = [
    'Confirmar firewall alcanzable en Sophos Central + S2S_Vaihingen Connection en rojo.',
    'Capturar logs IPsec antes de tocar nada (desaparecen al recuperar el túnel).',
    'Bounce del túnel: Active OFF → esperar ~10 s → ON. Renegocia Phase 1 + 2; vuelve verde en 30–60 s. Requiere "ok" explícito del cliente registrado en el ticket (regla N1: producción solo en P1 con autorización).',
    'Si no recupera → escalar N2 Sistemas, P1 con logs + estado de WAN.',
]
for item in runbook:
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_hr(doc)

# Notas de proceso
doc.add_heading('Notas de proceso', level=2)
notas = [
    'Clasificación y respuesta dentro del SLA de P1 (15 min).',
    'Cero modificaciones en producción desde CoolNetworks → sin riesgo añadido al incidente.',
    'Redirección de canal (WhatsApp → portal) realizada de forma proactiva, sin pedir permiso, manteniendo la sensación de acompañamiento.',
    'Conocimiento del entorno del cliente capturado en memoria para futuros incidentes del mismo patrón.',
]
for item in notas:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

out = r'g:\Mi unidad\Emprendimientos\AI_Automation_Agency\AIClieffNotes\WeeklyCompetions\week3\coolnetworks-tecnico-n1\coolnetworks-IT-Agent\reports\2026-06-08_maswer_vaihingen-internet-outage-s2s-vpn.docx'
doc.save(out)

import subprocess
subprocess.run(
    ['powershell', '-Command', f'Unblock-File -LiteralPath "{out}"'],
    check=True
)
print('OK')
